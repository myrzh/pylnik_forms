from zipfile import ZipFile
from csv import DictReader
from os import mkdir
from shutil import rmtree

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


GROUP_NUMBER = "5151003/30002"


def extract_csv(practice_number: str) -> None:
    with ZipFile(f"ORG_{practice_number}.csv.zip", "r") as zip_ref:
        zip_ref.extractall("responses_extracted")


def actualize_records(records: list) -> list:
    used_names = []
    responses = []

    for record in records[::-1]:
        if record["Имя"] not in used_names:
            responses.append(record)
            used_names.append(record["Имя"])
    responses.reverse()

    return responses


def init_document(practice_number: str) -> Document:
    document = Document()

    overall_style = document.styles["Normal"]
    overall_font = overall_style.font
    overall_font.name = "Times New Roman"
    overall_font.size = Pt(12)

    header_paragraph = document.add_paragraph(
        f"Практика {practice_number}, группа {GROUP_NUMBER}\n"
    )
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    return document


def write_response(document: Document, response: dict, is_last: bool) -> Document:
    name_paragraph = document.add_paragraph(response["Имя"])
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.style = document.styles["Heading 1"]

    for index in range(2, 4):
        document.add_paragraph(
            list(response.keys())[index].replace(" без заголовка", "")
        )
        document.add_paragraph(list(response.values())[index])

    if not is_last:
        document.add_page_break()

    return document


def main():
    practice_number = input("Enter practice number: ")

    extract_csv(practice_number)

    with open(
        f"responses_extracted/ORG_{practice_number}.csv", encoding="utf8"
    ) as csvfile:
        reader = DictReader(csvfile, delimiter=",", quotechar='"')
        responses = actualize_records(list(reader))
        responses.sort(key=lambda resp: resp["Имя"].lower())

    document = init_document(practice_number)

    for index, response in enumerate(responses):
        write_response(document, response, index == len(responses) - 1)

    try:
        mkdir("documents")
    except FileExistsError:
        pass
    document.save(f"documents/Практика {practice_number}.docx")
    rmtree("responses_extracted", ignore_errors=True)


if __name__ == "__main__":
    main()
